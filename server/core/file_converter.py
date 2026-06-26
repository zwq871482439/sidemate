# -*- coding: utf-8 -*-
"""
core/file_converter.py — 工作区文件格式转换（format_convert 工具的核心逻辑）

支持的转换路径（仅工作区内文件，绝不碰知识库）：
  md → docx   复用 doc_action.generate_docx
  md → txt    剥离 markdown 标记
  docx → md   python-docx 读段落，重建简易 markdown
  docx → txt  python-docx 读纯文本
  pdf → txt   pdfplumber 提取文本
  txt → md    原样存（纯文本视为 markdown 段落）

不支持的转换（超出范围，工具返回明确错误）：
  docx → pdf / pdf → docx / 任意 → pdf（需要排版引擎，风险高）
"""
import os
import re
import logging

log = logging.getLogger(__name__)

# 支持的格式
SUPPORTED_FORMATS = ("md", "docx", "txt", "pdf")
# 可作为"源"读取的格式（pdf 只能读不能写）
READABLE = {"md", "docx", "txt", "pdf"}
# 可作为"目标"写入的格式（pdf 不可写）
WRITABLE = {"md", "docx", "txt"}


def _strip_markdown(md_text: str) -> str:
    """把 markdown 剥离成纯文本（去标记，保留文字内容）。"""
    text = md_text
    # 代码块 → 保留内容
    text = re.sub(r"```[^\n]*\n(.*?)```", r"\1", text, flags=re.DOTALL)
    # 行内代码
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # 图片/链接 → 链接文字或 URL
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    # 加粗/斜体
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    # 标题标记
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # 列表标记
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    # 引用
    text = re.sub(r"^\s*>\s?", "", text, flags=re.MULTILINE)
    # 水平线
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)
    return text.strip()


def _read_docx(path: str) -> str:
    """读取 docx，返回 markdown 格式文本（标题/段落/列表简易重建）。"""
    from docx import Document
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            lines.append("")
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            # Heading 1 → #，Heading 2 → ##
            try:
                level = int(style.replace("heading", "").strip())
            except (ValueError, TypeError):
                level = 1
            level = max(1, min(6, level))
            lines.append("#" * level + " " + text)
        elif style == "list bullet":
            lines.append("- " + text)
        elif style.startswith("list"):
            lines.append("- " + text)
        else:
            lines.append(text)
    # 表格转 markdown
    for table in doc.tables:
        lines.append("")
        rows = table.rows
        if not rows:
            continue
        for ri, row in enumerate(rows):
            cells = [(c.text or "").strip().replace("|", "\\|") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if ri == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")
    return "\n".join(lines).strip()


def _read_pdf(path: str) -> str:
    """读取 pdf，返回纯文本（pdfplumber 提取）。"""
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                texts.append(t)
        return "\n\n".join(texts).strip()
    except ImportError:
        # fallback: pypdf
        from pypdf import PdfReader
        reader = PdfReader(path)
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        return "\n\n".join(texts).strip()


def _write_docx(md_text: str, path: str, title: str = "文档"):
    """md → docx，复用 doc_action.generate_docx。"""
    from pipelines.doc_action import generate_docx
    generate_docx(md_text, path, title=title)


def convert(src_path: str, dst_path: str) -> dict:
    """执行格式转换。

    Args:
        src_path: 源文件绝对路径（调用方负责校验在工作区内）
        dst_path: 目标文件绝对路径（调用方负责校验在工作区内）

    Returns:
        {"ok": bool, "src_format": str, "dst_format": str, "chars": int, "error": str}
    """
    if not os.path.exists(src_path):
        return {"ok": False, "error": "源文件不存在: %s" % os.path.basename(src_path)}

    src_ext = os.path.splitext(src_path)[1].lower().lstrip(".")
    dst_ext = os.path.splitext(dst_path)[1].lower().lstrip(".")

    if src_ext not in READABLE:
        return {"ok": False, "error": "不支持的源格式: .%s（支持 %s）" % (src_ext, "/".join(sorted(READABLE)))}
    if dst_ext not in WRITABLE:
        return {"ok": False, "error": "不支持的目标格式: .%s（支持 %s）" % (dst_ext, "/".join(sorted(WRITABLE)))}
    if src_ext == dst_ext and src_path == dst_path:
        return {"ok": False, "error": "源和目标是同一文件"}
    if src_ext == "pdf" and dst_ext != "txt":
        return {"ok": False, "error": "PDF 只能转为 txt（提取文本），不能转为 %s" % dst_ext}

    try:
        # 1. 读源文件 → 统一的中间 markdown 文本
        if src_ext == "md":
            with open(src_path, "r", encoding="utf-8") as f:
                content_md = f.read()
        elif src_ext == "txt":
            with open(src_path, "r", encoding="utf-8") as f:
                content_md = f.read()  # txt 视为纯文本 markdown
        elif src_ext == "docx":
            content_md = _read_docx(src_path)
        elif src_ext == "pdf":
            content_md = _read_pdf(src_path)
        else:
            return {"ok": False, "error": "未知源格式"}

        # 2. 中间文本 → 目标格式
        os.makedirs(os.path.dirname(dst_path), exist_ok=True)
        if dst_ext == "md":
            with open(dst_path, "w", encoding="utf-8") as f:
                f.write(content_md)
        elif dst_ext == "txt":
            with open(dst_path, "w", encoding="utf-8") as f:
                f.write(_strip_markdown(content_md))
        elif dst_ext == "docx":
            title = os.path.splitext(os.path.basename(dst_path))[0]
            _write_docx(content_md, dst_path, title=title)
        else:
            return {"ok": False, "error": "未知目标格式"}

        chars = len(content_md)
        log.info("[CONVERT] %s → %s (%d字)", src_ext, dst_ext, chars)
        return {
            "ok": True,
            "src_format": src_ext,
            "dst_format": dst_ext,
            "chars": chars,
            "dst_name": os.path.basename(dst_path),
        }
    except Exception as e:
        log.error("[CONVERT] 转换失败 %s→%s: %s", src_ext, dst_ext, str(e)[:150])
        return {"ok": False, "error": "转换失败: %s" % str(e)[:150]}

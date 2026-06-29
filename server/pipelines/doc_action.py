# -*- coding: utf-8 -*-
"""
doc_action.py — 文档生成 Action（Patch1 V2 — 两阶段）
=====================================================
两阶段流程：
  Phase 1（提纲）：KB搜索 → 生成提纲 → yield doc_outline → 结束
  Phase 2（正文）：前端带 doc_continue 参数 → 基于提纲生成完整文档 → yield 正文

用户确认流程由前端控制：
  - Phase 1 结束后前端展示提纲 + 确认/取消按钮
  - 用户点确认 → 前端发新请求带 doc_continue=<outline>
  - 用户点取消 → 前端不做任何操作

调用方式：chat.py 的 doc 分支 for 循环 yield run_doc_action() 的产出
"""

import os
import re
import time
import logging
import threading

log = logging.getLogger(__name__)


# 全局 KB 上下文缓存（Phase 1 → Phase 2 传递）
_kb_context_cache = {}
_kb_cache_lock = threading.Lock()


def run_doc_action(
    message: str,
    mgr,           # ModelManager
    model_name: str,
    max_tokens: int = None,
    history: list = None,
    kb=None,       # KnowledgeBase 实例
    context_cache: dict = None,
    strategy_enhancement: str = "",
    doc_continue: str = "",  # Phase 2: 用户确认的提纲内容
    kb_doc_content: str = "",  # 用户引用的 KB 文档全文
):
    """
    执行文档生成 Action（两阶段）。

    Phase 1 (doc_continue 为空):
      1. KB 搜索
      2. 模型生成提纲
      3. yield ("doc_outline", outline_text) → 结束

    Phase 2 (doc_continue 不为空):
      1. 基于提纲 + KB 上下文生成完整文档
      2. yield 正文 token 流

    Yields:
        (phase, content) — 兼容 chat_stream 接口
    """
    # ====== Phase 2: 基于已确认提纲生成完整文档 ======
    if doc_continue:
        yield from _run_phase2(
            message, mgr, model_name, max_tokens, history,
            context_cache, strategy_enhancement, doc_continue,
            kb_doc_content=kb_doc_content
        )
        return

    # ====== Phase 1: 生成提纲 ======
    yield from _run_phase1(
        message, mgr, model_name, max_tokens, history,
        kb, context_cache, strategy_enhancement,
        kb_doc_content=kb_doc_content
    )


def _run_phase1(
    message, mgr, model_name, max_tokens, history,
    kb, context_cache, strategy_enhancement,
    kb_doc_content=""
):
    """Phase 1: KB搜索 → 生成提纲 → yield doc_outline"""

    # Step 1: KB 自动搜索（修 #doc模式搜索永远失败：KnowledgeBase 无 query 方法，改用 get_context）
    kb_context = ""
    kb_refs = []
    if kb:
        try:
            kb_loaded = getattr(kb, '_embedder_loaded', False)
            if kb_loaded:
                yield ("mode_hint", " 正在搜索文库...")
                kb_context, kb_refs = kb.get_context(message, max_chars=2000, ai_mode="local")
                if kb_refs:
                    # 列出检索到的来源（与 chat 模式 _base.py:250 对齐）
                    _src_labels = []
                    for ref in kb_refs[:3]:
                        _lbl = ref.get("source_label") or ref.get("filename") or ref.get("doc_id") or "?"
                        _src_labels.append(_lbl)
                    yield ("mode_hint", " 已检索文库（%d 条相关文档：%s），正在生成文档提纲..." % (
                        len(kb_refs), "、".join(_src_labels)))
                    # 发送 kb_sources 事件，让前端渲染来源卡片（与 chat 模式一致）
                    yield ("kb_sources", [{"label": s.get("source_label", "?"),
                                           "snippet": (s.get("text_snippet", "") or s.get("content", ""))[:100]}
                                          for s in kb_refs[:5]])
                    log.info("[DOC_ACTION] KB 搜索到 %d 条参考资料: %s" % (len(kb_refs), "、".join(_src_labels)))
                else:
                    yield ("mode_hint", " 文库中未找到相关内容，将直接生成文档提纲。")
                    log.info("[DOC_ACTION] KB 无检索结果，fallback 直接生成提纲")
        except Exception as e:
            log.warning("[DOC_ACTION] KB 搜索失败: %s" % str(e)[:100])

    # 缓存 KB 上下文（Phase 2 会用到）
    with _kb_cache_lock:
        _kb_context_cache["last"] = kb_context
        _kb_context_cache["refs_count"] = len(kb_refs)

    # Step 2: 生成提纲
    yield ("mode_hint", " 正在生成文档提纲...")

    from prompts import DOC_OUTLINE_PROMPT
    outline_prompt = DOC_OUTLINE_PROMPT.format(user_request=message)

    # 注入用户引用的 KB 文档全文（最高优先级，放在 prompt 开头）
    if kb_doc_content:
        outline_prompt = "[用户引用了文库文档，内容如下：]\n%s\n\n%s" % (kb_doc_content, outline_prompt)
        log.info("[DOC_ACTION] 用户引用 KB 文档注入 Phase1: %d 字" % len(kb_doc_content))

    if kb_context:
        outline_prompt = "%s\n\n[以下是文库检索到的参考资料：]\n%s\n\n请结合以上参考资料生成提纲。" % (
            outline_prompt, kb_context)

    # 收集提纲文本
    outline_text = ""
    for phase, content in mgr.chat_stream(
        outline_prompt, model_name, max_tokens or 1024, history,
        context_cache=context_cache,
        strategy_enhancement=strategy_enhancement,
    ):
        if phase == "text":
            outline_text += content
            yield ("text", content)  # 流式展示提纲
        elif phase == "raw":
            outline_text += content
            yield ("text", content)
        elif phase == "fold":
            yield ("fold", content)
        elif phase == "task_type":
            yield (phase, content)

    if not outline_text.strip():
        yield ("mode_hint", "提纲生成失败，请重试。")
        return

    # Step 3: yield 提纲确认事件（前端会展示确认/取消按钮）
    log.info("[DOC_ACTION] 提纲生成完成 (%d 字)" % len(outline_text))
    yield ("doc_outline", outline_text.strip())


def _run_phase2(
    message, mgr, model_name, max_tokens, history,
    context_cache, strategy_enhancement, outline,
    kb_doc_content=""
):
    """Phase 2: 基于已确认的提纲生成完整文档"""

    yield ("mode_hint", " 正在基于提纲生成完整文档...")

    from prompts import DOC_FULL_PROMPT

    # 获取 Phase 1 缓存的 KB 上下文
    with _kb_cache_lock:
        kb_context = _kb_context_cache.pop("last", "")

    full_prompt = DOC_FULL_PROMPT.format(user_request=message, outline=outline)

    # 注入用户引用的 KB 文档全文（Phase 2 也需要）
    if kb_doc_content:
        full_prompt = "[用户引用了文库文档，内容如下：]\n%s\n\n%s" % (kb_doc_content, full_prompt)
        log.info("[DOC_ACTION] 用户引用 KB 文档注入 Phase2: %d 字" % len(kb_doc_content))

    if kb_context:
        full_prompt = "%s\n\n[以下是可参考的文库资料：]\n%s" % (full_prompt, kb_context)

    # 流式生成正文
    for phase, content in mgr.chat_stream(
        full_prompt, model_name, max_tokens, history,
        context_cache=context_cache,
        strategy_enhancement=strategy_enhancement,
    ):
        yield (phase, content)

    log.info("[DOC_ACTION] 完整文档生成完成")


# ============================================================
#  Markdown → .docx 转换
# ============================================================

def _parse_markdown_to_sections(md_text: str):
    """将 Markdown 文本解析为 (title, sections) 结构

    Returns:
        (title, [(heading, body), ...])
    """
    lines = md_text.strip().split('\n')
    title = "文档"
    sections = []
    current_heading = None
    current_body_lines = []

    for line in lines:
        if line.startswith('# ') and not line.startswith('## '):
            title = line.lstrip('# ').strip()
            continue
        # P6：支持 ### / #### 子标题（映射到 Heading 2/3）
        if re.match(r'^##\s', line) and not re.match(r'^###\s', line):
            if current_heading is not None or current_body_lines:
                sections.append((current_heading or "概述", '\n'.join(current_body_lines).strip()))
            current_heading = line.lstrip('# ').strip()
            current_body_lines = []
            continue
        if re.match(r'^###\s', line):
            if current_heading is not None or current_body_lines:
                sections.append((current_heading or "概述", '\n'.join(current_body_lines).strip()))
            current_heading = "  " + line.lstrip('# ').strip()  # 缩进表示子标题
            current_body_lines = []
            continue
        current_body_lines.append(line)

    if current_heading is not None or current_body_lines:
        body = '\n'.join(current_body_lines).strip()
        if body:
            # Patch5 修复：没标题就留空，不强加"总结"
            # 让模型自己决定要不要加标题
            sections.append((current_heading or "", body))

    if not sections and md_text.strip():
        # 兜底：纯文本无标题，整个作为正文
        sections.append(("", md_text.strip()))

    return title, sections


def generate_docx(content: str, output_path: str, title: str = "文档"):
    """使用 pandoc 将 Markdown 转为 .docx，再统一字体为等线

    三层兜底（修 #5-b/#5-d：pypandoc 未安装时整个函数崩溃）：
      1. pypandoc（若已安装）—— 首选
      2. subprocess 直调 pandoc CLI —— pypandoc 缺失时用，pandoc 二进制通常存在
      3. _generate_docx_manual —— pandoc 也不可用时的纯 python-docx 回退
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    _FONT_CN = '等线'
    _FONT_EN = 'Calibri'

    # 用 pandoc 转换 markdown → docx
    # 注意：不传 --metadata title，因为模型正文里的 # 一级标题就是文档标题，
    # 传 metadata title 会导致标题区出现两层（pandoc 元数据标题 + 正文 H1）
    _pandoc_ok = False
    # 兜底1：pypandoc（import 移进 try，避免 ModuleNotFoundError 冒泡导致无法回退）
    try:
        import pypandoc
        pypandoc.convert_text(
            content, 'docx', format='markdown',
            outputfile=output_path,
            extra_args=[
                '--from=markdown+autolink_bare_uris+task_lists',
            ]
        )
        _pandoc_ok = True
    except ImportError:
        log.info("[DOC] pypandoc 未安装，尝试 subprocess 直调 pandoc CLI")
    except Exception as e:
        log.warning("[DOC] pypandoc 转换失败，尝试 CLI: %s", e)

    # 兜底2：subprocess 直调 pandoc CLI（pypandoc 缺失或失败时）
    if not _pandoc_ok:
        try:
            import subprocess
            proc = subprocess.run(
                ["pandoc", "--from=markdown+autolink_bare_uris+task_lists",
                 "--to=docx", "-o", output_path],
                input=content, encoding="utf-8",
                capture_output=True, timeout=60,
            )
            if proc.returncode == 0 and os.path.exists(output_path):
                _pandoc_ok = True
            else:
                log.warning("[DOC] pandoc CLI 失败 rc=%s: %s",
                            proc.returncode, (proc.stderr or "")[:200])
        except FileNotFoundError:
            log.warning("[DOC] pandoc 二进制未找到，回退手动生成")
        except Exception as e:
            log.warning("[DOC] pandoc CLI 异常，回退手动生成: %s", e)

    # 兜底3：pandoc 全失败 → 纯 python-docx 手动生成
    if not _pandoc_ok:
        return _generate_docx_manual(content, output_path, title)

    # 打开生成的 docx，统一字体
    doc = Document(output_path)

    # 覆盖 docDefaults 字体
    from lxml import etree
    _W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _doc_defaults = doc.styles.element.find('{%s}docDefaults' % _W_NS)
    if _doc_defaults is not None:
        _rpr_default = _doc_defaults.find('{%s}rPrDefault' % _W_NS)
        if _rpr_default is None:
            _rpr_default = etree.SubElement(_doc_defaults, '{%s}rPrDefault' % _W_NS)
        _rpr = _rpr_default.find('{%s}rPr' % _W_NS)
        if _rpr is None:
            _rpr = etree.SubElement(_rpr_default, '{%s}rPr' % _W_NS)
        _rfonts = _rpr.find('{%s}rFonts' % _W_NS)
        if _rfonts is None:
            _rfonts = etree.SubElement(_rpr, '{%s}rFonts' % _W_NS)
        _rfonts.set('{%s}ascii' % _W_NS, _FONT_EN)
        _rfonts.set('{%s}hAnsi' % _W_NS, _FONT_EN)
        _rfonts.set('{%s}eastAsia' % _W_NS, _FONT_CN)
        _rfonts.set('{%s}cs' % _W_NS, _FONT_EN)
        for _attr in ['asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme']:
            _full = '{%s}%s' % (_W_NS, _attr)
            if _full in _rfonts.attrib:
                del _rfonts.attrib[_full]

    # 遍历所有段落，统一字体 + 排版美化
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_LINE_SPACING
    _TITLE_COLOR = RGBColor(0x1F, 0x29, 0x37)    # 深灰，标题色
    _ACCENT_COLOR = RGBColor(0x2D, 0x4A, 0x6F)   # 品牌蓝，H1
    for _para in doc.paragraphs:
        _style_name = _para.style.name if _para.style else ''
        # 统一字体
        for _run in _para.runs:
            _run.font.name = _FONT_EN
            _run._element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)
        # 排版美化：按样式类型设置字号/间距/颜色
        if _style_name == 'Title' or _style_name == 'Heading 1':
            for _run in _para.runs:
                _run.font.size = Pt(20)
                _run.font.bold = True
                _run.font.color.rgb = _ACCENT_COLOR
            _para.paragraph_format.space_before = Pt(18)
            _para.paragraph_format.space_after = Pt(12)
        elif _style_name == 'Heading 2':
            for _run in _para.runs:
                _run.font.size = Pt(15)
                _run.font.bold = True
                _run.font.color.rgb = _TITLE_COLOR
            _para.paragraph_format.space_before = Pt(14)
            _para.paragraph_format.space_after = Pt(8)
        elif _style_name == 'Heading 3':
            for _run in _para.runs:
                _run.font.size = Pt(13)
                _run.font.bold = True
                _run.font.color.rgb = _TITLE_COLOR
            _para.paragraph_format.space_before = Pt(10)
            _para.paragraph_format.space_after = Pt(6)
        else:
            # 正文段落：统一行距 + 字号
            for _run in _para.runs:
                if not _run.font.size:
                    _run.font.size = Pt(11)
            _para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            _para.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    log.info("[DOC] .docx 生成完成 (pandoc): %s (%d bytes)", os.path.basename(output_path), file_size)


# ===== HTML 可视化报告（含 mermaid 图表，自包含单文件） =====

_mermaid_js_cache = None  # mermaid.min.js 全文缓存，避免每次读盘


def _load_mermaid_js() -> str:
    """读取 mermaid.min.js 全文（带缓存）"""
    global _mermaid_js_cache
    if _mermaid_js_cache is not None:
        return _mermaid_js_cache
    # mermaid.min.js 在 server/static/vendor/
    _here = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    candidates = [
        os.path.join(_here, "static", "vendor", "mermaid.min.js"),
        os.path.join(_here, "static", "vendor", "mermaid.js"),
    ]
    for p in candidates:
        if os.path.isfile(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    _mermaid_js_cache = f.read()
                log.info("[DOC] mermaid.min.js 已加载 (%d bytes)", len(_mermaid_js_cache))
                return _mermaid_js_cache
            except Exception as e:
                log.warning("[DOC] 读取 mermaid.min.js 失败: %s", e)
    log.warning("[DOC] 未找到 mermaid.min.js，HTML 报告图表将无法渲染")
    _mermaid_js_cache = ""
    return _mermaid_js_cache


_marked_js_cache = None


def _load_marked_js() -> str:
    """读取 marked.min.js 全文（带缓存）——用于报告内前端解析 LLM 混用的 Markdown 语法"""
    global _marked_js_cache
    if _marked_js_cache is not None:
        return _marked_js_cache
    _here = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    p = os.path.join(_here, "static", "vendor", "marked.min.js")
    try:
        with open(p, "r", encoding="utf-8") as f:
            _marked_js_cache = f.read()
        log.info("[DOC] marked.min.js 已加载 (%d bytes)", len(_marked_js_cache))
    except Exception as e:
        log.warning("[DOC] 读取 marked.min.js 失败: %s", e)
        _marked_js_cache = ""
    return _marked_js_cache


# HTML 报告 CSS：专业报告排版（参考高端数据报告设计）+ mermaid 交互 + 打印优化
# 设计理念：暖灰背景(护眼)、衬线标题(高级感)、语义色、柔和阴影、大圆角
# LLM 可用预设 class，也可自己写 <style> 覆盖
_HTML_REPORT_CSS = """
* { box-sizing: border-box; margin: 0; padding: 0; }
:root {
  --c-primary: #2f6f5e; --c-primary-light: #4a8e7c; --c-accent: #c97b3f;
  --c-text: #14171e; --c-text-soft: #3a4256; --c-muted: #7c7a72;
  --c-border: #e9e3d6; --c-border-soft: #f0ebde; --c-bg: #f7f5f0;
  --c-surface: #ffffff; --c-card: #ffffff;
  --c-good: #2f6f5e; --c-bad: #b06367; --c-warn: #c97b3f;
  --radius: 14px; --radius-sm: 8px;
  --shadow: 0 2px 8px rgba(0,0,0,.03); --shadow-lg: 0 4px 20px rgba(0,0,0,.06);
  --serif: Georgia, "Times New Roman", "Noto Serif SC", serif;
  --sans: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", Roboto, sans-serif;
  --mono: "JetBrains Mono", "Cascadia Code", Consolas, monospace;
}
body {
  font-family: var(--sans); background: var(--c-bg); color: var(--c-text);
  line-height: 1.6; font-size: 15px;
}
.wrap { max-width: 900px; margin: 0 auto; padding: 48px 36px 80px; }
h1 { font-family: var(--serif); font-weight: 400; font-size: 34px; margin-bottom: 8px; letter-spacing: -.01em; }
h2 { font-family: var(--serif); font-weight: 400; font-size: 23px; margin: 44px 0 14px; padding-bottom: 8px; border-bottom: 1px solid var(--c-border); }
h2 .accent { color: var(--c-primary); }
h3 { font-size: 17px; font-weight: 600; margin: 24px 0 8px; color: var(--c-text); }
h4 { font-size: 14px; font-weight: 600; margin: 16px 0 6px; color: var(--c-text-soft); text-transform: uppercase; letter-spacing: .06em; }
p { margin: 10px 0; color: var(--c-text-soft); }
a { color: var(--c-primary); text-decoration: none; }
a:hover { text-decoration: underline; }
hr { border: none; border-top: 1px solid var(--c-border); margin: 32px 0; }

/* 副标题/元信息 */
.subtitle { color: var(--c-muted); font-size: 14.5px; margin-bottom: 4px; }
.meta { color: var(--c-muted); font-size: 13px; margin-bottom: 28px; padding-bottom: 20px; border-bottom: 1px solid var(--c-border); }

/* 表格 */
table { width: 100%; border-collapse: collapse; font-size: 13.5px; margin: 14px 0; }
thead { background: var(--c-border-soft); }
th, td { text-align: left; padding: 9px 13px; border-bottom: 1px solid var(--c-border); }
th { font-weight: 600; color: var(--c-text-soft); font-size: 11.5px; text-transform: uppercase; letter-spacing: .05em; }
tr:hover { background: var(--c-border-soft); }
.pos { color: var(--c-good); font-weight: 500; }
.neg { color: var(--c-bad); font-weight: 500; }

/* 代码 */
code { font-family: var(--mono); font-size: 12.5px; background: var(--c-border-soft); padding: 1px 6px; border-radius: 3px; }
pre { background: #1a1a2e; color: #e2e8f0; padding: 16px 20px; border-radius: var(--radius-sm); overflow-x: auto; line-height: 1.5; }
pre code { background: transparent; color: inherit; padding: 0; font-size: 13px; }
ul, ol { padding-left: 22px; margin: 8px 0; } li { margin: 5px 0; }
img { max-width: 100%; border-radius: var(--radius-sm); }

/* ===== 高级组件库（参考专业数据报告设计）===== */

/* 导读块 lead：渐变背景突出核心结论 */
.lead { background: linear-gradient(135deg, rgba(201,123,63,.08), rgba(47,111,94,.06)); border-radius: var(--radius); padding: 24px 28px; margin: 20px 0; font-size: 15px; line-height: 1.7; color: var(--c-text-soft); }
.lead strong { color: var(--c-text); }

/* 大号数字 highlight-num：视觉焦点 */
.highlight-num { font-family: var(--serif); font-size: 22px; color: var(--c-accent); margin: 0 3px; font-weight: 400; }

/* 提示框 callout：分色左边框（比 note 更精致） */
.callout { background: rgba(47,111,94,.06); border-left: 3px solid var(--c-primary); padding: 14px 18px; border-radius: 0 var(--radius-sm) var(--radius-sm) 0; margin: 16px 0; font-size: 14px; color: var(--c-text-soft); }
.callout.warn { background: rgba(201,123,63,.07); border-left-color: var(--c-warn); }
.callout.danger { background: rgba(176,99,103,.07); border-left-color: var(--c-bad); }
.callout.tip { background: rgba(74,142,124,.07); border-left-color: var(--c-primary-light); }
.callout strong { color: var(--c-text); }

/* 兼容旧 class（note/tip/warning/danger/success → callout 风格）*/
.note, .tip, .warning, .danger, .success {
  padding: 14px 18px; margin: 16px 0; border-radius: 0 var(--radius-sm) var(--radius-sm) 0;
  border-left: 3px solid; font-size: 14px; color: var(--c-text-soft);
}
.note { background: rgba(47,111,94,.06); border-color: var(--c-primary); }
.tip { background: rgba(74,142,124,.07); border-color: var(--c-primary-light); }
.warning { background: rgba(201,123,63,.07); border-color: var(--c-warn); }
.danger { background: rgba(176,99,103,.07); border-color: var(--c-bad); }
.success { background: rgba(47,111,94,.07); border-color: var(--c-good); }

/* 卡片 */
.card { background: var(--c-surface); border: 1px solid var(--c-border); border-radius: var(--radius); padding: 20px 24px; margin: 16px 0; box-shadow: var(--shadow); }
.card-title { font-weight: 600; font-size: 15px; margin-bottom: 8px; color: var(--c-text); }

/* 网格 */
.grid-2 { display: grid; grid-template-columns: 1fr 1fr; gap: 18px; margin: 18px 0; }
.grid-3 { display: grid; grid-template-columns: repeat(3,1fr); gap: 16px; margin: 18px 0; }
@media (max-width: 640px) { .grid-2, .grid-3 { grid-template-columns: 1fr; } }

/* 统计数字块 stat：大号衬线数字 */
.stats { display: flex; gap: 18px; flex-wrap: wrap; margin: 18px 0; }
.stat { flex: 1; min-width: 140px; }
.stat.card { text-align: center; }
.stat h4 { font-size: 11px; color: var(--c-muted); font-weight: 500; text-transform: uppercase; letter-spacing: .08em; margin-bottom: 8px; }
.stat-num { font-family: var(--serif); font-size: 28px; line-height: 1.1; color: var(--c-primary); font-weight: 400; }
.stat-num.pos { color: var(--c-good); }
.stat-num.neg { color: var(--c-bad); }
.stat-num.accent { color: var(--c-accent); }
.stat-label { font-size: 12px; color: var(--c-muted); margin-top: 5px; }

/* 标签 */
.badge { display: inline-block; padding: 2px 10px; border-radius: 999px; font-size: 11.5px; font-weight: 500; background: rgba(47,111,94,.1); color: var(--c-primary); }
.badge.green { background: rgba(5,150,105,.12); color: #065f46; }
.badge.orange { background: rgba(245,158,11,.12); color: #92400e; }
.badge.red { background: rgba(176,99,103,.12); color: var(--c-bad); }
.badge.gray { background: var(--c-border-soft); color: var(--c-muted); }

/* 进度条 */
.progress { background: var(--c-border); border-radius: 999px; height: 8px; overflow: hidden; margin: 8px 0; }
.progress-bar { background: var(--c-primary); height: 100%; border-radius: 999px; transition: width .3s ease; }

/* 时间线 */
.timeline { border-left: 2px solid var(--c-border); padding-left: 20px; margin: 16px 0; }
.timeline-item { margin: 14px 0; position: relative; }
.timeline-item::before { content: ''; position: absolute; left: -26px; top: 6px; width: 10px; height: 10px; border-radius: 50%; background: var(--c-accent); }

/* 高亮文本 */
.highlight { background: linear-gradient(transparent 60%, rgba(201,123,63,.25) 60%); padding: 0 3px; }

/* 概念解释块 concept */
.concept { background: rgba(74,142,124,.06); border-left: 3px solid var(--c-primary-light); padding: 12px 18px; border-radius: 0 var(--radius-sm) var(--radius-sm) 0; margin: 12px 0; font-size: 14px; color: var(--c-text-soft); }
.concept strong { color: var(--c-text); }

/* 页脚 */
.footer { margin-top: 56px; padding-top: 20px; border-top: 1px solid var(--c-border); color: var(--c-muted); font-size: 12.5px; text-align: center; }

/* 现代风格（LLM 加 class=vibrant 到外层 div 启用） */
.vibrant { background: linear-gradient(135deg, #f0f4ff 0%, #fdf4ff 100%); border-radius: var(--radius); padding: 32px; }
.vibrant h1 { background: linear-gradient(135deg, #2f6f5e, #8b5cf6); -webkit-background-clip: text; background-clip: text; color: transparent; }
.vibrant .card { border: none; box-shadow: var(--shadow-lg); }
.vibrant .stat { background: rgba(255,255,255,.7); backdrop-filter: blur(4px); }

/* ===== mermaid 图表交互容器 ===== */
.chart-frame { border: 1px solid var(--c-border); border-radius: var(--radius); overflow: hidden; margin: 18px 0; position: relative; background: var(--c-surface); }
.cf-hint { position: absolute; top: 6px; right: 8px; font-size: 11px; color: var(--c-muted); background: rgba(255,255,255,.85); padding: 2px 8px; border-radius: 4px; z-index: 5; pointer-events: none; }
.chart-stage { overflow: hidden; cursor: grab; padding: 16px 8px; min-height: 80px; text-align: center; }
.chart-stage:active { cursor: grabbing; }
.chart-stage svg { max-width: none !important; height: auto !important; transition: none; display: inline-block; }
.chart-toolbar { position: absolute; bottom: 6px; right: 6px; display: flex; gap: 2px; align-items: center; background: rgba(255,255,255,.9); border: 1px solid var(--c-border); border-radius: 6px; padding: 2px 4px; font-size: 11px; z-index: 5; opacity: 0; transition: opacity .15s; }
.chart-frame:hover .chart-toolbar { opacity: 1; }
.chart-toolbar button { border: none; background: transparent; cursor: pointer; width: 22px; height: 22px; font-size: 14px; border-radius: 4px; color: var(--c-text); line-height: 1; }
.chart-toolbar button:hover { background: var(--c-border); }
.chart-toolbar .zoom-val { min-width: 36px; text-align: center; color: var(--c-muted); font-variant-numeric: tabular-nums; }

/* ===== 顶部提示条（可关闭）===== */
.report-tipbar { position: sticky; top: 0; z-index: 50; background: linear-gradient(135deg, var(--c-primary), #6366f1); color: #fff; padding: 8px 20px; font-size: 13px; display: flex; align-items: center; justify-content: center; gap: 6px; margin: -48px -36px 24px; border-radius: 0; }
.report-tipbar .tipbar-close { position: absolute; right: 14px; top: 50%; transform: translateY(-50%); cursor: pointer; opacity: .8; font-size: 16px; }
.report-tipbar .tipbar-close:hover { opacity: 1; }

/* 打印优化 */
@media print {
  body { background: #fff; }
  .wrap { max-width: none; padding: 0; }
  .report-tipbar, .chart-toolbar { display: none !important; }
  .chart-stage { overflow: visible !important; }
  .chart-stage svg { transform: none !important; }
  pre, .chart-frame, .card, table { page-break-inside: avoid; }
  h1, h2, h3 { page-break-after: avoid; }
}
"""


def _convert_markdown_in_html(content):
    """轻量 Markdown→HTML 转换（兜底，当 LLM 在 HTML body 里混用 Markdown 时）

    只处理最常见的：标题(#/##/###/####)、表格(|...|)、引用(>)、粗体(**)、
    无序列表(-/*)、有序列表(1.)、分隔线(---)。
    已是 HTML 标签的行原样保留（检测 < 开头）。
    ```mermaid``` 围栏保护，不转换。
    """
    # 先保护 mermaid 围栏块（占位符，后面还原）
    placeholders = []
    def _protect(m):
        placeholders.append(m.group(0))
        return "\x00MERMAID%d\x00" % (len(placeholders) - 1)
    content = re.sub(r"```mermaid\s*\n.*?```", _protect, content, flags=re.DOTALL)

    # 保护已有 HTML 块标签（<div <table <h1 <ul 等），整行以 < 开头的不转
    lines = content.split("\n")
    out = []
    in_table = False
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 还原 mermaid 占位符（单独成行时）
        if stripped.startswith("\x00MERMAID") and stripped.endswith("\x00"):
            idx = int(stripped[8:-1])
            out.append(placeholders[idx])
            i += 1
            continue

        # 已是 HTML 标签开头 → 原样保留（含 <div <h1 <p <table <section 等）
        if re.match(r"\s*<\w", line):
            out.append(line)
            i += 1
            continue

        # 空行
        if not stripped:
            if in_table:
                out.append("</tbody></table>")
                in_table = False
            out.append("")
            i += 1
            continue

        # 分隔线 ---
        if re.match(r"^-{3,}$", stripped):
            out.append("<hr>")
            i += 1
            continue

        # 标题 #/##/###/####
        m = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            out.append("<h%d>%s</h%d>" % (level, _md_inline(m.group(2)), level))
            i += 1
            continue

        # 引用 >
        if stripped.startswith(">"):
            quote_text = _md_inline(stripped[1:].strip())
            out.append("<blockquote>%s</blockquote>" % quote_text)
            i += 1
            continue

        # 表格 | ... |
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # 分隔行 |---|---|
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                i += 1
                continue
            if not in_table:
                out.append("<table><thead><tr>" + "".join("<th>%s</th>" % _md_inline(c) for c in cells) + "</tr></thead><tbody>")
                in_table = True
            else:
                out.append("<tr>" + "".join("<td>%s</td>" % _md_inline(c) for c in cells) + "</tr>")
            i += 1
            continue

        # 无序列表 - /*
        if re.match(r"^[-*]\s+", stripped):
            out.append("<li>%s</li>" % _md_inline(re.sub(r"^[-*]\s+", "", stripped)))
            i += 1
            continue

        # 有序列表 1.
        m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if m:
            out.append("<li>%s</li>" % _md_inline(m.group(1)))
            i += 1
            continue

        # 普通段落
        out.append("<p>%s</p>" % _md_inline(stripped))
        i += 1

    if in_table:
        out.append("</tbody></table>")

    result = "\n".join(out)
    # 还原行内的 mermaid 占位符
    for idx, original in enumerate(placeholders):
        result = result.replace("\x00MERMAID%d\x00" % idx, original)
    return result


def _md_inline(text):
    """行内 Markdown 转换：**粗体** → <strong>，`代码` → <code>"""
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    return text


def generate_html_report(content: str, output_path: str, title: str = "报告"):
    """生成自包含的 HTML 可视化报告（内联 mermaid.js + 交互，单文件可独立打开）

    - ```mermaid``` 围栏自动转成可缩放/拖拽的交互容器
    - 内联 mermaid.min.js + 交互增强脚本
    - 顶部提示条（滚轮缩放/Ctrl+P 转 PDF，可关闭）
    - 预设 CSS 组件库（note/card/grid/stats/badge 等），LLM 可直接用 class

    Args:
        content: LLM 写的 HTML body 内容（可含 ```mermaid``` 围栏代码块）
        output_path: 输出 .html 文件路径
        title: 报告标题
    """
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    mermaid_js = _load_mermaid_js()
    marked_js = _load_marked_js()

    # P7: mermaid 围栏用占位符保护（marked 会把 ``` 围栏转成 <pre><code>，必须先抽离）
    import json as _json
    _mermaid_codes = []
    _frame_counter = [0]

    def _fence_to_frame(m):
        code = m.group(1).strip()
        idx = _frame_counter[0]
        _frame_counter[0] += 1
        _mermaid_codes.append(code)
        return (
            '<div class="chart-frame">'
            '<div class="cf-hint">滚轮缩放 · 拖拽移动 · 双击复位</div>'
            '<div class="chart-stage" data-stage>'
            '<div class="chart-slot" data-idx="%d"></div>' % idx +
            '</div>'
            '<div class="chart-toolbar">'
            '<button type="button" data-act="out" title="缩小">−</button>'
            '<span class="zoom-val">100%</span>'
            '<button type="button" data-act="in" title="放大">+</button>'
            '<button type="button" data-act="reset" title="复位">⟲</button>'
            '</div>'
            '</div>'
        )

    content = re.sub(
        r"```mermaid\s*\n.*?```",
        _fence_to_frame,
        content,
        flags=re.DOTALL,
    )

    # P7: marked 在前端处理 LLM 混用的 Markdown 语法。
    # 把 LLM 内容（已抽取 mermaid）JSON 序列化塞进 body，前端 marked 解析。
    # marked 对已是 HTML 的内容原样保留（不重解析），只把 Markdown 转 HTML——解决标签错乱。
    # marked min js 39KB，与 mermaid min js 3MB 一起内联到 HTML（gzip 后总 ~1MB），可接受。
    _md_input = _json.dumps(content, ensure_ascii=False)
    processed_content = (
        '<div id="md-src" data-md=' + _json.dumps(_md_input, ensure_ascii=False) + ' style="display:none"></div>'
        '<script>setTimeout(function(){var s=document.getElementById("md-src");'
        'if(s&&typeof marked==="function"){s.outerHTML=marked.parse(s.dataset.md);}},0);</script>'
    )

    # 源码数组（JSON 序列化，安全转义）
    _codes_json = _json.dumps(_mermaid_codes, ensure_ascii=False)

    # 顶部提示条
    tipbar = (
        '<div class="report-tipbar" id="reportTipbar">'
        '💡 图表支持滚轮缩放/拖拽。需要 PDF/打印？按 <b>Ctrl+P</b> 另存'
        '<span class="tipbar-close" onclick="var t=document.getElementById(\'reportTipbar\');'
        'if(t)t.style.display=\'none\';try{localStorage.setItem(\'reportTipHidden\',\'1\')}catch(e){}">×</span>'
        '</div>'
    )

    # 交互增强脚本（原生 JS，不依赖任何库；验证过的稳定方案）
    # mermaid 源码从注入的 JS 数组读取（不在 DOM 暴露 mermaid 文本）
    interact_js = """
<script>
var _MERMAID_CODES = %s;  // 由 Python 注入的源码数组
function initReportMermaid(){
  if(typeof mermaid==='undefined'){setTimeout(initReportMermaid,50);return;}
  mermaid.initialize({startOnLoad:false,theme:'default',securityLevel:'loose',
    flowchart:{useMaxWidth:false,padding:12}});
  var slots=document.querySelectorAll('.chart-slot[data-idx]');
  slots.forEach(function(slot){
    var idx=parseInt(slot.getAttribute('data-idx')||'0',10);
    var code=_MERMAID_CODES[idx];
    if(!code){enhanceMermaidStage(slot.closest('[data-stage]'));return;}
    var id='m'+idx+'-'+Math.random().toString(36).slice(2,6);
    mermaid.render(id,code).then(function(res){
      slot.innerHTML=res.svg;
      slot.className='chart-rendered';
      enhanceMermaidStage(slot.closest('[data-stage]'));
    }).catch(function(e){
      slot.innerHTML='<pre style="color:#dc2626;font-size:12px">图表渲染失败: '+String(e.message||e).slice(0,150)+'</pre>';
      enhanceMermaidStage(slot.closest('[data-stage]'));
    });
  });
  try{if(localStorage.getItem('reportTipHidden')==='1'){var t=document.getElementById('reportTipbar');if(t)t.style.display='none';}}catch(e){}
}
function enhanceMermaidStage(stage){
  if(!stage||stage._enhanced)return;
  stage._enhanced=true;
  var svg=stage.querySelector('svg');if(!svg)return;
  var frame=stage.closest('.chart-frame');
  var zoomVal=frame?frame.querySelector('.zoom-val'):null;
  var toolbar=frame?frame.querySelector('.chart-toolbar'):null;
  var scale=1,tx=0,ty=0,MIN=0.3,MAX=3;
  function apply(){svg.style.transform='translate('+tx+'px,'+ty+'px) scale('+scale+')';svg.style.transformOrigin='center center';if(zoomVal)zoomVal.textContent=Math.round(scale*100)+'%%';}
  function reset(){scale=1;tx=0;ty=0;apply();}
  stage.addEventListener('wheel',function(e){e.preventDefault();var d=e.deltaY<0?0.1:-0.1;scale=Math.max(MIN,Math.min(MAX,scale+d));apply();},{passive:false});
  var drag=false,sx,sy,ox,oy;
  stage.addEventListener('mousedown',function(e){drag=true;sx=e.clientX;sy=e.clientY;ox=tx;oy=ty;e.preventDefault();});
  document.addEventListener('mousemove',function(e){if(!drag)return;tx=ox+(e.clientX-sx);ty=oy+(e.clientY-sy);apply();});
  document.addEventListener('mouseup',function(){drag=false;});
  stage.addEventListener('dblclick',reset);
  if(toolbar){toolbar.addEventListener('click',function(e){var b=e.target.closest('button');if(!b)return;var a=b.getAttribute('data-act');if(a==='in'){scale=Math.min(MAX,scale+0.2);apply();}else if(a==='out'){scale=Math.max(MIN,scale-0.2);apply();}else if(a==='reset'){reset();}});}
}
initReportMermaid();
</script>
""" % _codes_json

    html = (
        '<!DOCTYPE html>\n<html lang="zh-CN">\n<head>\n'
        '<meta charset="utf-8">\n'
        '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
        '<title>' + title.replace("<", "&lt;").replace(">", "&gt;") + '</title>\n'
        '<style>' + _HTML_REPORT_CSS + '</style>\n'
        '</head>\n<body>\n'
    )
    html += tipbar + '\n'
    html += processed_content
    # mermaid.min.js 和 marked.min.js 都放 body 末尾（content 之后）：
    # 若放 head，库会在解析到 .mermaid/.md 元素时自动处理破坏 DOM
    # 放 body 末尾 + 初始化脚本紧随其后，元素先存在再被我们手动处理
    if marked_js:
        html += '\n<script>' + marked_js + '</script>\n'
    if mermaid_js:
        html += '\n<script>' + mermaid_js + '</script>\n'
    html += '\n' + (interact_js if mermaid_js else "") + '\n'
    html += '</body>\n</html>\n'

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)

    file_size = os.path.getsize(output_path)
    log.info("[DOC] .html 报告生成完成: %s (%d bytes, mermaid=%s, frames=%d)",
             os.path.basename(output_path), file_size, "yes" if mermaid_js else "no", _frame_counter[0])


# ===== PPT 演示文稿（reveal.js，含 mermaid 图，自包含单文件） =====

_reveal_js_cache = None
_reveal_css_cache = None
_reveal_theme_cache = None


def _load_reveal_js():
    """读取 reveal.min.js + reveal.css + 主题 CSS（带缓存，全部内联）"""
    global _reveal_js_cache, _reveal_css_cache, _reveal_theme_cache
    if _reveal_js_cache is not None:
        return _reveal_js_cache, _reveal_css_cache, _reveal_theme_cache
    _here = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    vdir = os.path.join(_here, "static", "vendor")
    def _read(name):
        p = os.path.join(vdir, name)
        if os.path.isfile(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception:
                pass
        return ""
    _reveal_js_cache = _read("reveal.min.js")
    _reveal_css_cache = _read("reveal.css")
    _reveal_theme_cache = _read("reveal-theme-black.css")  # 默认黑底白字
    log.info("[DOC] reveal.js 资源已加载 (js=%d css=%d theme=%d)",
             len(_reveal_js_cache), len(_reveal_css_cache), len(_reveal_theme_cache))
    return _reveal_js_cache, _reveal_css_cache, _reveal_theme_cache


# PPT slide 内的辅助样式（补充 reveal 主题，让 mermaid/表格/列表更好看）
_PPT_SLIDE_CSS = """
/* PPT slide 内容样式 — 专业美观，信息密度合理
   关键：所有字号用 rem（不随 reveal 基础字号 42px 放大），rem 根 = 16px 可控
   之前用 em 在 reveal 里被放大约 2.6 倍 → 字号爆炸溢出
   reveal slide 尺寸 ~960×700（16:9），内容要排得下 */
:root { font-size: 16px; }  /* 显式锁住 rem 根 */
.reveal section { text-align: left; padding: 0.4rem 1rem; font-size: 1rem; }  /* 16px 不放大 */
.reveal h1 { font-size: 2.2rem; line-height: 1.2; margin-bottom: 0.3rem; }
.reveal h2 { font-size: 1.5rem; line-height: 1.25; margin: 0.4rem 0 0.3rem; }
.reveal h3 { font-size: 1.1rem; color: #a8b2d1; margin-bottom: 0.25rem; }
.reveal p { font-size: 0.9rem; line-height: 1.5; margin: 0.3rem 0; }
.reveal ul, .reveal ol { font-size: 0.85rem; line-height: 1.5; margin-left: 1.2rem; }
.reveal li { margin: 0.15rem 0; }
.reveal strong { color: #64ffda; }

/* 表格 */
.reveal table { font-size: 0.75rem; border-collapse: collapse; margin: 0.4rem auto; width: 92%; }
.reveal th { background: rgba(100,255,218,.12); border-bottom: 2px solid #64ffda; padding: 5px 10px; font-weight: 600; }
.reveal td { border-bottom: 1px solid #333; padding: 4px 10px; }
.reveal tr:nth-child(even) { background: rgba(255,255,255,.03); }

/* 代码 */
.reveal pre { font-size: 0.65rem; margin: 0.3rem auto; width: 92%; }
.reveal code { font-size: 0.95em; }

/* 卡片 */
.reveal .card { background: rgba(255,255,255,.06); border: 1px solid rgba(255,255,255,.1); border-radius: 8px; padding: 10px 12px; margin: 0.3em 0; }
.reveal .card-title { font-size: 0.95rem; font-weight: 600; color: #64ffda; margin-bottom: 4px; }

/* 网格 */
.reveal .grid-2, .reveal .grid-3 { gap: 10px; margin: 0.3em 0; }
.reveal .grid-2 { display: grid; grid-template-columns: 1fr 1fr; }
.reveal .grid-3 { display: grid; grid-template-columns: repeat(3,1fr); }

/* 数据统计块 */
.reveal .stats { display: flex; gap: 14px; margin: 0.4rem 0; justify-content: center; }
.reveal .stat { flex: 1; text-align: center; padding: 8px; }
.reveal .stat-num { font-size: 1.6rem; font-weight: 700; color: #64ffda; line-height: 1.1; }
.reveal .stat-label { font-size: 0.7rem; color: #8892b0; margin-top: 2px; }

/* 标签 */
.reveal .badge { font-size: 0.7rem; padding: 2px 8px; border-radius: 999px; background: rgba(100,255,218,.15); color: #64ffda; }
.reveal .badge.green { background: rgba(5,150,105,.2); color: #6ee7b7; }
.reveal .badge.orange { background: rgba(245,158,11,.2); color: #fcd34d; }

/* mermaid 图 */
.reveal .chart-frame { background: #fff; color: #333; border-radius: 10px; margin: 0.4em auto; padding: 12px; box-shadow: 0 4px 20px rgba(0,0,0,.3); }
.reveal .chart-stage { padding: 8px; min-height: 60px; text-align: center; }
.reveal .chart-stage svg { max-width: 88% !important; height: auto !important; }
.reveal .cf-hint, .reveal .chart-toolbar { display: none; }

/* 引用 */
.reveal blockquote { font-size: 0.85rem; border-left: 3px solid #64ffda; padding: 8px 14px; margin: 0.3rem 0; background: rgba(100,255,218,.05); font-style: italic; }

/* 高亮 */
.reveal .highlight { background: rgba(100,255,218,.25); padding: 0 4px; border-radius: 3px; }

/* 进度条 */
.reveal .progress { background: rgba(100,255,218,.15); }

/* 提示条 */
.ppt-tipbar { position: fixed; top: 0; left: 0; right: 0; z-index: 100; background: rgba(59,130,246,.95);
  color: #fff; padding: 6px 20px; font-size: 13px; text-align: center; display: flex; align-items: center;
  justify-content: center; gap: 8px; }
.ppt-tipbar .close { position: absolute; right: 14px; cursor: pointer; opacity: .8; font-size: 16px; }
"""


def generate_ppt_html(content: str, output_path: str, title: str = "演示文稿"):
    """生成自包含的 PPT 演示文稿（内联 reveal.js + mermaid.js，单文件可独立演示）

    - LLM 写的内容应是多个 `<section>`（每张幻灯片），系统包进 reveal.js 结构
    - ```mermaid``` 围栏自动转成可交互容器（复用报告的 mermaid 注入逻辑）
    - 用户浏览器打开 → 方向键翻页 → 按 F 全屏 → 加 ?print-pdf 导出 PDF

    Args:
        content: LLM 写的幻灯片内容（多个 <section>，可含 ```mermaid``` 围栏）
        output_path: 输出 .ppt.html 文件路径
        title: 演示文稿标题
    """
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    mermaid_js = _load_mermaid_js()
    reveal_js, reveal_css, reveal_theme = _load_reveal_js()

    # 把 ```mermaid``` 围栏转成占位容器（复用报告的逻辑，源码进 JS 数组）
    import json as _json
    _mermaid_codes = []
    _frame_counter = [0]

    def _fence_to_frame(m):
        code = m.group(1).strip()
        idx = _frame_counter[0]
        _frame_counter[0] += 1
        _mermaid_codes.append(code)
        return (
            '<div class="chart-frame"><div class="chart-stage" data-stage>'
            '<div class="chart-slot" data-idx="%d"></div></div></div>' % idx
        )

    processed_content = re.sub(
        r"```mermaid\s*\n(.*?)```",
        _fence_to_frame,
        content,
        flags=re.DOTALL,
    )
    _codes_json = _json.dumps(_mermaid_codes, ensure_ascii=False)

    # 顶部提示条
    tipbar = (
        '<div class="ppt-tipbar" id="pptTipbar">'
        '⬅ ➡ 方向键翻页 · 按 <b>F</b> 全屏 · 导出 PDF 请在网址后加 <b>?print-pdf</b>'
        '<span class="close" onclick="var t=document.getElementById(\'pptTipbar\');'
        'if(t)t.style.display=\'none\';try{localStorage.setItem(\'pptTipHidden\',\'1\')}catch(e){}">×</span>'
        '</div>'
    )

    # PPT 的 mermaid 渲染脚本（slidechanged 时渲染当前页的图，解决懒加载时序）
    ppt_mermaid_js = """
<script>
var _MERMAID_CODES = %s;
function initPptMermaid(){
  if(typeof mermaid==='undefined'){setTimeout(initPptMermaid,50);return;}
  mermaid.initialize({startOnLoad:false,theme:'dark',securityLevel:'loose',
    flowchart:{useMaxWidth:false,padding:12}});
  function renderSlide(slide){
    if(!slide)return;
    var slots=slide.querySelectorAll('.chart-slot[data-idx]:not([data-done])');
    slots.forEach(function(slot){
      var idx=parseInt(slot.getAttribute('data-idx')||'0',10);
      var code=_MERMAID_CODES[idx];if(!code)return;
      slot.setAttribute('data-done','1');
      var id='p'+idx+'-'+Math.random().toString(36).slice(2,6);
      mermaid.render(id,code).then(function(res){slot.innerHTML=res.svg;})
        .catch(function(e){slot.innerHTML='<pre style="color:#f66;font-size:11px">图表渲染失败</pre>';});
    });
  }
  // 初始化时渲染当前页
  if(typeof Reveal!=='undefined'){
    renderSlide(Reveal.getCurrentSlide());
    Reveal.addEventListener('slidechanged',function(ev){renderSlide(ev.currentSlide);});
    Reveal.addEventListener('ready',function(ev){renderSlide(ev.currentSlide);});
  }else{
    // Reveal 未就绪，轮询
    var t=setInterval(function(){
      if(typeof Reveal!=='undefined'&&Reveal.getCurrentSlide){
        clearInterval(t);renderSlide(Reveal.getCurrentSlide());
        Reveal.addEventListener('slidechanged',function(ev){renderSlide(ev.currentSlide);});
      }
    },100);
  }
  try{if(localStorage.getItem('pptTipHidden')==='1'){var t=document.getElementById('pptTipbar');if(t)t.style.display='none';}}catch(e){}
}
</script>
""" % _codes_json

    html = (
        '<!DOCTYPE html>\n<html lang="zh-CN">\n<head>\n'
        '<meta charset="utf-8">\n'
        '<meta name="viewport" content="width=device-width, initial-scale=1, maximum-scale=1, user-scalable=no">\n'
        '<title>' + title.replace("<", "&lt;").replace(">", "&gt;") + '</title>\n'
        '<style>' + reveal_css + '</style>\n'
        '<style>' + reveal_theme + '</style>\n'
        '<style>' + _PPT_SLIDE_CSS + '</style>\n'
        '</head>\n<body>\n'
        + tipbar + '\n'
        '<div class="reveal">\n<div class="slides">\n'
        + processed_content + '\n'
        '</div>\n</div>\n'
    )
    # reveal.js 必须先加载并 initialize，mermaid 在其后
    if reveal_js:
        html += '<script>' + reveal_js + '</script>\n'
        html += '<script>Reveal.initialize({controls:true,progress:true,center:true,'
        html += 'hash:true,slideNumber:true,viewDistance:5,transition:"slide"});</script>\n'
    if mermaid_js:
        html += '<script>' + mermaid_js + '</script>\n'
    html += ppt_mermaid_js
    html += '<script>initPptMermaid();</script>\n'
    html += '</body>\n</html>\n'

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)

    file_size = os.path.getsize(output_path)
    log.info("[DOC] .ppt.html 演示文稿生成完成: %s (%d bytes, slides mermaid=%d)",
             os.path.basename(output_path), file_size, _frame_counter[0])


def _generate_docx_manual(content: str, output_path: str, title: str = "文档"):
    """pandoc 不可用时的手动回退方案"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc_title, sections = _parse_markdown_to_sections(content)
    doc = Document()

    _FONT_CN = '等线'
    _FONT_EN = 'Calibri'
    _FONT_SIZE = Pt(11)

    from lxml import etree
    _W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _doc_defaults = doc.styles.element.find('{%s}docDefaults' % _W_NS)
    if _doc_defaults is not None:
        _rpr_default = _doc_defaults.find('{%s}rPrDefault' % _W_NS)
        if _rpr_default is None:
            _rpr_default = etree.SubElement(_doc_defaults, '{%s}rPrDefault' % _W_NS)
        _rpr = _rpr_default.find('{%s}rPr' % _W_NS)
        if _rpr is None:
            _rpr = etree.SubElement(_rpr_default, '{%s}rPr' % _W_NS)
        _rfonts = _rpr.find('{%s}rFonts' % _W_NS)
        if _rfonts is None:
            _rfonts = etree.SubElement(_rpr, '{%s}rFonts' % _W_NS)
        _rfonts.set('{%s}ascii' % _W_NS, _FONT_EN)
        _rfonts.set('{%s}hAnsi' % _W_NS, _FONT_EN)
        _rfonts.set('{%s}eastAsia' % _W_NS, _FONT_CN)
        _rfonts.set('{%s}cs' % _W_NS, _FONT_EN)
        for _attr in ['asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme']:
            _full = '{%s}%s' % (_W_NS, _attr)
            if _full in _rfonts.attrib:
                del _rfonts.attrib[_full]

    style = doc.styles['Normal']
    style.font.name = _FONT_EN
    style.font.size = _FONT_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)
    _title_style = doc.styles['Title']
    _title_style.font.name = _FONT_EN
    _title_style.element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)
    for level in range(1, 5):
        hstyle = doc.styles['Heading %d' % level]
        hstyle.font.name = _FONT_EN
        hstyle.element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)

    title_para = doc.add_heading(doc_title or title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _run in title_para.runs:
        _run.font.name = _FONT_EN
        _run.font.size = Pt(14)
        _run._element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)

    for heading, body in sections:
        if heading:
            _h_level = 2 if heading.startswith('  ') else 1
            _h_text = heading.lstrip() if _h_level == 2 else heading
            _h_para = doc.add_heading(_h_text, level=_h_level)
            for _run in _h_para.runs:
                _run.font.name = _FONT_EN
                _run._element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)
        if not body:
            continue
        paragraphs = re.split(r'\n{2,}', body)
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            list_items = []
            _list_has_marks = False
            for line in para_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                line = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', line)
                line = re.sub(r'!\[([^\]]*)\]\([^)]*\)', r'\1', line)
                if re.match(r'^\d+[.、．)\s]', line):
                    _list_has_marks = True
                    list_items.append(line)
                elif re.match(r'^[-*•]\s', line):
                    _list_has_marks = True
                    list_items.append('• ' + re.sub(r'^[-*•]\s+', '', line))
                else:
                    if _list_has_marks and list_items:
                        list_items.append(line)
                    else:
                        list_items.append(line)
            for item in list_items:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', item)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = _FONT_SIZE
                        run.font.name = _FONT_EN
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.size = _FONT_SIZE
                        run.font.name = 'Consolas'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)
                    elif part:
                        run = p.add_run(part)
                        run.font.size = _FONT_SIZE
                        run.font.name = _FONT_EN
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), _FONT_CN)

    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    log.info("[DOC] .docx 生成完成 (手动): %s (%d bytes)", os.path.basename(output_path), file_size)

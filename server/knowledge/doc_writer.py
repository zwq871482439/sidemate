# -*- coding: utf-8 -*-
"""
doc_writer.py - Word 文档生成技能

模板驱动生成 .docx 文件:
  - 报告模板: 标题 + 正文段落（支持标题层级）
  - 摘要模板: 简短标题 + 总结段落
  - 信函模板: 收件人 + 正文 + 落款
  - 自由格式: 自定义段落列表

输出到 files/ 沙箱目录。

依赖: python-docx
"""
__version__ = "v1.0"

import os, sys, json, logging
from typing import Optional, List, Dict

log = logging.getLogger(__name__)

# ===== 常量 =====
SANDBOX_ENV_KEY = "_sandbox_dir"
WORKSPACE_ENV_KEY = "_workspace_dir"
MAX_FILENAME_LEN = 128                # 文件名最大长度
MAX_CONTENT_TOTAL_CHARS = 200000       # 文档内容总量上限
DEFAULT_FONT_NAME = "Microsoft YaHei"  # 默认中文字体
DEFAULT_FONT_SIZE = 11                 # 默认正文字号 (pt)
DEFAULT_HEADING_FONT_SIZE = 16         # 默认标题字号 (pt)
FILENAME_PATTERN = r'^[\w\-. \u4e00-\u9fff]+$'

# 模板类型
TPL_REPORT = "report"
TPL_SUMMARY = "summary"
TPL_LETTER = "letter"
TPL_CUSTOM = "custom"

ALLOWED_TEMPLATES = (TPL_REPORT, TPL_SUMMARY, TPL_LETTER, TPL_CUSTOM)


def _load_params() -> dict:
    """从 stdin 读取 JSON 参数"""
    try:
        raw = sys.stdin.read()
        return json.loads(raw) if raw.strip() else {}
    except (json.JSONDecodeError, IOError) as e:
        return {"error": "参数读取失败: %s" % str(e)}


def _validate_path(path: str, sandbox_dir: str) -> Optional[str]:
    """校验路径"""
    if not path:
        return None
    sandbox_abs = os.path.realpath(sandbox_dir)
    target_abs = os.path.realpath(os.path.join(sandbox_dir, path))
    if not target_abs.startswith(sandbox_abs + os.sep) and target_abs != sandbox_abs:
        return None
    return target_abs


def _safe_filename(name: str) -> bool:
    """检查文件名是否安全"""
    import re
    return bool(re.match(FILENAME_PATTERN, name)) and len(name) <= MAX_FILENAME_LEN


def _add_paragraph_with_font(doc, text: str, style=None, font_size: int = None,
                              bold: bool = False, font_name: str = DEFAULT_FONT_NAME):
    """添加段落并设置字体"""
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = font_name
        if font_size:
            run.font.size = __import__('docx.shared', fromlist=['Pt']).Pt(font_size)
        run.font.bold = bold
    return para


def _build_report(params: dict) -> dict:
    """生成报告模板

    参数:
      title: 报告标题
      author: 作者（可选）
      date: 日期（可选，默认今天）
      sections: 段落列表 [{"heading": "xxx", "level": 1, "content": "xxx"}, ...]
      filename: 输出文件名
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return {"error": "缺少依赖 python-docx，请运行: pip install python-docx"}

    title = params.get("title", "未命名报告")
    author = params.get("author", "")
    date_str = params.get("date", __import__("datetime").datetime.now().strftime("%Y-%m-%d"))
    sections = params.get("sections", [])

    doc = Document()

    # 标题
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.name = DEFAULT_FONT_NAME

    # 元信息
    meta_parts = []
    if author:
        meta_parts.append("作者: %s" % author)
    meta_parts.append("日期: %s" % date_str)
    doc.add_paragraph(" | ".join(meta_parts))

    # 分隔线
    doc.add_paragraph("---")

    # 正文段落
    for sec in sections:
        heading = sec.get("heading", "")
        level = min(int(sec.get("level", 1)), 4)
        content = sec.get("content", "")

        if heading:
            h = doc.add_heading(heading, level=level)
            for run in h.runs:
                run.font.name = DEFAULT_FONT_NAME

        if content:
            para = doc.add_paragraph(content)
            for run in para.runs:
                run.font.name = DEFAULT_FONT_NAME
                run.font.size = Pt(DEFAULT_FONT_SIZE)

    return {"doc": doc}


def _build_summary(params: dict) -> dict:
    """生成摘要模板

    参数:
      title: 摘要标题
      summary: 摘要内容
      key_points: 要点列表 ["xxx", "xxx"]
      filename: 输出文件名
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return {"error": "缺少依赖 python-docx，请运行: pip install python-docx"}

    title = params.get("title", "摘要")
    summary_text = params.get("summary", "")
    key_points = params.get("key_points", [])

    doc = Document()

    # 标题
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.name = DEFAULT_FONT_NAME

    # 摘要正文
    if summary_text:
        para = doc.add_paragraph(summary_text)
        for run in para.runs:
            run.font.name = DEFAULT_FONT_NAME
            run.font.size = Pt(DEFAULT_FONT_SIZE)

    # 要点
    if key_points:
        doc.add_heading("要点", level=2)
        for point in key_points:
            para = doc.add_paragraph(point, style="List Bullet")
            for run in para.runs:
                run.font.name = DEFAULT_FONT_NAME
                run.font.size = Pt(DEFAULT_FONT_SIZE)

    return {"doc": doc}


def _build_letter(params: dict) -> dict:
    """生成信函模板

    参数:
      recipient: 收件人
      subject: 主题
      body: 正文内容
      sender: 发件人（可选）
      filename: 输出文件名
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return {"error": "缺少依赖 python-docx，请运行: pip install python-docx"}

    recipient = params.get("recipient", "")
    subject = params.get("subject", "无主题")
    body = params.get("body", "")
    sender = params.get("sender", "")

    doc = Document()

    # 收件人
    if recipient:
        para = doc.add_paragraph("致: %s" % recipient)
        for run in para.runs:
            run.font.name = DEFAULT_FONT_NAME
            run.font.size = Pt(DEFAULT_FONT_SIZE)

    doc.add_paragraph("")

    # 主题
    h = doc.add_heading(subject, level=1)
    for run in h.runs:
        run.font.name = DEFAULT_FONT_NAME

    # 正文
    if body:
        para = doc.add_paragraph(body)
        for run in para.runs:
            run.font.name = DEFAULT_FONT_NAME
            run.font.size = Pt(DEFAULT_FONT_SIZE)

    # 落款
    doc.add_paragraph("")
    if sender:
        para = doc.add_paragraph("此致\n%s" % sender)
        for run in para.runs:
            run.font.name = DEFAULT_FONT_NAME

    return {"doc": doc}


def _build_custom(params: dict) -> dict:
    """生成自由格式文档

    参数:
      paragraphs: 段落列表 [{"text": "xxx", "style": "heading/normal/bullet", "level": 1}, ...]
      filename: 输出文件名
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return {"error": "缺少依赖 python-docx，请运行: pip install python-docx"}

    paragraphs = params.get("paragraphs", [])

    doc = Document()

    for p in paragraphs:
        text = p.get("text", "")
        style = p.get("style", "normal")
        level = min(int(p.get("level", 1)), 4)

        if style == "heading":
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = DEFAULT_FONT_NAME
        elif style == "bullet":
            para = doc.add_paragraph(text, style="List Bullet")
            for run in para.runs:
                run.font.name = DEFAULT_FONT_NAME
                run.font.size = Pt(DEFAULT_FONT_SIZE)
        else:
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.name = DEFAULT_FONT_NAME
                run.font.size = Pt(DEFAULT_FONT_SIZE)

    return {"doc": doc}


# ===== 模板分发 =====

_BUILDERS = {
    TPL_REPORT: _build_report,
    TPL_SUMMARY: _build_summary,
    TPL_LETTER: _build_letter,
    TPL_CUSTOM: _build_custom,
}


def main():
    """技能入口"""
    params = _load_params()

    if "error" in params:
        print(json.dumps({"status": "error", "error": params["error"]}, ensure_ascii=False))
        return

    sandbox_dir = params.get(SANDBOX_ENV_KEY, "")
    if not sandbox_dir:
        print(json.dumps({"status": "error", "error": "沙箱目录未指定"}, ensure_ascii=False))
        return

    os.makedirs(sandbox_dir, exist_ok=True)

    # 获取模板类型
    template = params.get("template", TPL_REPORT)
    if template not in _BUILDERS:
        print(json.dumps({
            "status": "error",
            "error": "未知模板: %s (允许: %s)" % (template, ", ".join(_BUILDERS.keys())),
        }, ensure_ascii=False))
        return

    # 获取输出文件名
    filename = params.get("filename", "output.docx")
    if not filename.endswith(".docx"):
        filename += ".docx"

    basename = os.path.basename(filename)
    if not _safe_filename(basename):
        print(json.dumps({"status": "error", "error": "文件名不安全: %s" % basename}, ensure_ascii=False))
        return

    output_path = os.path.join(sandbox_dir, filename)
    output_abs = os.path.realpath(output_path)
    sandbox_abs = os.path.realpath(sandbox_dir)
    if not output_abs.startswith(sandbox_abs + os.sep):
        print(json.dumps({"status": "error", "error": "非法输出路径"}, ensure_ascii=False))
        return

    # 构建
    builder = _BUILDERS[template]
    result = builder(params)

    if "error" in result:
        print(json.dumps({"status": "error", "error": result["error"]}, ensure_ascii=False))
        return

    # 保存
    doc = result["doc"]
    try:
        doc.save(output_path)
    except Exception as e:
        print(json.dumps({"status": "error", "error": "保存失败: %s" % str(e)[:100]}, ensure_ascii=False))
        return

    file_size = os.path.getsize(output_path)

    print(json.dumps({
        "status": "ok",
        "filename": filename,
        "path": filename,
        "size": file_size,
        "template": template,
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()

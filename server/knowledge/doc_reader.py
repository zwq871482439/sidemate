# -*- coding: utf-8 -*-
"""
doc_reader.py - Word 文档读取技能

分段读取 .docx 文件:
  - 首次读取: 前 HEAD_MAX_CHARS 字 + 尾部 TAIL_MAX_CHARS 字 + 中间段落标题
  - 表格提取: 结构信息（行列数 + 表头）
  - 追问读取: 指定段落/页码的完整内容

依赖: python-docx
"""
__version__ = "v1.0"

import os, sys, json, logging
from typing import Optional, List, Dict

log = logging.getLogger(__name__)

# ===== 常量 =====
SANDBOX_ENV_KEY = "_sandbox_dir"
WORKSPACE_ENV_KEY = "_workspace_dir"
HEAD_MAX_CHARS = 2000          # 首段最大字符数
TAIL_MAX_CHARS = 500           # 尾段最大字符数
TITLE_MAX_CHARS = 100          # 每个标题摘要最大字符数
MAX_TABLE_ROWS_PREVIEW = 5     # 表格预览行数
MAX_FULL_CONTENT_CHARS = 50000 # 完整读取最大字符数

# 读取模式
MODE_PREVIEW = "preview"       # 首次读取（首尾+标题摘要）
MODE_FULL = "full"             # 完整读取
MODE_SECTION = "section"       # 指定段落/标题
MODE_TABLE = "table"           # 指定表格完整内容

# 标题样式关键词
HEADING_KEYWORDS = ("Heading", "Title", "heading", "title", "TOC")


def _load_params() -> dict:
    """从 stdin 读取 JSON 参数"""
    try:
        raw = sys.stdin.read()
        return json.loads(raw) if raw.strip() else {}
    except (json.JSONDecodeError, IOError) as e:
        return {"error": "参数读取失败: %s" % str(e)}


def _validate_path(path: str, sandbox_dir: str) -> Optional[str]:
    """校验路径"""
    if not path:
        return None
    sandbox_abs = os.path.realpath(sandbox_dir)
    target_abs = os.path.realpath(os.path.join(sandbox_dir, path))
    if not target_abs.startswith(sandbox_abs + os.sep) and target_abs != sandbox_abs:
        return None
    return target_abs


def _is_heading(paragraph) -> bool:
    """判断段落是否是标题"""
    style_name = paragraph.style.name if paragraph.style else ""
    return any(kw in style_name for kw in HEADING_KEYWORDS)


def _get_heading_level(paragraph) -> int:
    """获取标题级别（1-6）"""
    style_name = paragraph.style.name if paragraph.style else ""
    import re
    m = re.search(r'(\d+)', style_name)
    if m:
        return min(int(m.group(1)), 6)
    if "Title" in style_name:
        return 1
    return 0


def _extract_paragraphs(doc) -> List[Dict]:
    """提取文档所有段落的结构化信息"""
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        is_head = _is_heading(para)
        level = _get_heading_level(para) if is_head else 0
        paragraphs.append({
            "index": i,
            "text": text,
            "is_heading": is_head,
            "heading_level": level,
            "style": para.style.name if para.style else "",
            "chars": len(text),
        })
    return paragraphs


def _extract_tables(doc) -> List[Dict]:
    """提取文档所有表格的结构信息"""
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = table.rows
        cols = table.columns
        row_count = len(rows)
        col_count = len(cols)

        # 提取表头（第一行）
        headers = []
        if row_count > 0:
            for cell in rows[0].cells:
                headers.append(cell.text.strip())

        # 预览前几行
        preview = []
        for row in rows[:MAX_TABLE_ROWS_PREVIEW]:
            row_data = [cell.text.strip() for cell in row.cells]
            preview.append(row_data)

        tables.append({
            "index": t_idx,
            "rows": row_count,
            "cols": col_count,
            "headers": headers,
            "preview": preview,
        })
    return tables


def _read_preview(file_path: str) -> dict:
    """首次预览读取: 首段 + 标题摘要 + 尾段 + 表格结构"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "缺少依赖 python-docx，请运行: pip install python-docx"}

    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": "文件读取失败: %s" % str(e)[:100]}

    paragraphs = _extract_paragraphs(doc)
    tables = _extract_tables(doc)

    if not paragraphs:
        return {
            "status": "ok",
            "mode": "preview",
            "file": os.path.basename(file_path),
            "total_paragraphs": 0,
            "head_text": "",
            "tail_text": "",
            "headings": [],
            "tables": tables,
            "total_chars": 0,
        }

    # 首段文本
    head_chars = 0
    head_parts = []
    for p in paragraphs:
        if head_chars + p["chars"] > HEAD_MAX_CHARS:
            remaining = HEAD_MAX_CHARS - head_chars
            head_parts.append(p["text"][:remaining] + "...")
            head_chars += remaining
            break
        head_parts.append(p["text"])
        head_chars += p["chars"]
    head_text = "\n".join(head_parts)

    # 标题摘要
    headings = []
    for p in paragraphs:
        if p["is_heading"]:
            headings.append({
                "level": p["heading_level"],
                "text": p["text"][:TITLE_MAX_CHARS],
                "index": p["index"],
            })

    # 尾段文本
    tail_chars = 0
    tail_parts = []
    for p in reversed(paragraphs):
        if tail_chars + p["chars"] > TAIL_MAX_CHARS:
            remaining = TAIL_MAX_CHARS - tail_chars
            tail_parts.insert(0, "..." + p["text"][-remaining:])
            tail_chars += remaining
            break
        tail_parts.insert(0, p["text"])
        tail_chars += p["chars"]
    tail_text = "\n".join(tail_parts)

    total_chars = sum(p["chars"] for p in paragraphs)

    return {
        "status": "ok",
        "mode": "preview",
        "file": os.path.basename(file_path),
        "total_paragraphs": len(paragraphs),
        "total_chars": total_chars,
        "head_text": head_text,
        "head_chars": head_chars,
        "tail_text": tail_text,
        "tail_chars": tail_chars,
        "headings": headings,
        "heading_count": len(headings),
        "tables": tables,
        "table_count": len(tables),
    }


def _read_full(file_path: str, max_chars: int = MAX_FULL_CONTENT_CHARS) -> dict:
    """完整读取文档"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "缺少依赖 python-docx，请运行: pip install python-docx"}

    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": "文件读取失败: %s" % str(e)[:100]}

    paragraphs = _extract_paragraphs(doc)
    tables = _extract_tables(doc)

    # 拼接全部段落
    all_text_parts = []
    total = 0
    for p in paragraphs:
        if total + p["chars"] > max_chars:
            remaining = max_chars - total
            all_text_parts.append(p["text"][:remaining] + "\n... (已截断)")
            total += remaining
            break
        prefix = "#" * p["heading_level"] + " " if p["is_heading"] else ""
        all_text_parts.append(prefix + p["text"])
        total += p["chars"]

    full_text = "\n".join(all_text_parts)

    # 完整表格内容
    full_tables = []
    for t_idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            rows_data.append([cell.text.strip() for cell in row.cells])
        full_tables.append({
            "index": t_idx,
            "rows": rows_data,
        })

    return {
        "status": "ok",
        "mode": "full",
        "file": os.path.basename(file_path),
        "content": full_text,
        "chars": total,
        "truncated": total >= max_chars,
        "total_paragraphs": len(paragraphs),
        "total_chars": sum(p["chars"] for p in paragraphs),
        "tables_full": full_tables,
        "table_count": len(full_tables),
    }


def _read_section(file_path: str, section_index: int = 0, section_title: str = "") -> dict:
    """读取指定段落/标题下的内容"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "缺少依赖 python-docx，请运行: pip install python-docx"}

    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": "文件读取失败: %s" % str(e)[:100]}

    paragraphs = _extract_paragraphs(doc)

    if section_title:
        # 按标题查找
        found_idx = None
        for i, p in enumerate(paragraphs):
            if p["is_heading"] and section_title.lower() in p["text"].lower():
                found_idx = i
                break
        if found_idx is None:
            return {"error": "未找到标题: %s" % section_title}
        section_index = found_idx

    if section_index < 0 or section_index >= len(paragraphs):
        return {"error": "段落索引越界: %d (总段落数: %d)" % (section_index, len(paragraphs))}

    # 从指定段落开始，读到下一个同级/更高级标题为止
    start_level = paragraphs[section_index]["heading_level"]
    section_parts = [paragraphs[section_index]["text"]]
    for i in range(section_index + 1, len(paragraphs)):
        p = paragraphs[i]
        if p["is_heading"] and p["heading_level"] <= start_level and start_level > 0:
            break
        section_parts.append(p["text"])

    section_text = "\n".join(section_parts)

    return {
        "status": "ok",
        "mode": "section",
        "file": os.path.basename(file_path),
        "section_index": section_index,
        "content": section_text,
        "chars": len(section_text),
        "paragraphs_in_section": len(section_parts),
    }


def _read_table(file_path: str, table_index: int = 0) -> dict:
    """读取指定表格的完整内容"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "缺少依赖 python-docx，请运行: pip install python-docx"}

    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": "文件读取失败: %s" % str(e)[:100]}

    if table_index < 0 or table_index >= len(doc.tables):
        return {"error": "表格索引越界: %d (总表格数: %d)" % (table_index, len(doc.tables))}

    table = doc.tables[table_index]
    rows_data = []
    for row in table.rows:
        rows_data.append([cell.text.strip() for cell in row.cells])

    return {
        "status": "ok",
        "mode": "table",
        "file": os.path.basename(file_path),
        "table_index": table_index,
        "rows": rows_data,
        "row_count": len(rows_data),
        "col_count": len(rows_data[0]) if rows_data else 0,
    }


# ===== 操作分发 =====

_READ_MODES = {
    MODE_PREVIEW: lambda p, fp: _read_preview(fp),
    MODE_FULL: lambda p, fp: _read_full(fp, int(p.get("max_chars", MAX_FULL_CONTENT_CHARS))),
    MODE_SECTION: lambda p, fp: _read_section(
        fp,
        section_index=int(p.get("section_index", 0)),
        section_title=p.get("section_title", ""),
    ),
    MODE_TABLE: lambda p, fp: _read_table(
        fp,
        table_index=int(p.get("table_index", 0)),
    ),
}


def main():
    """技能入口"""
    params = _load_params()

    if "error" in params:
        print(json.dumps({"status": "error", "error": params["error"]}, ensure_ascii=False))
        return

    sandbox_dir = params.get(SANDBOX_ENV_KEY, "")
    if not sandbox_dir:
        print(json.dumps({"status": "error", "error": "沙箱目录未指定"}, ensure_ascii=False))
        return

    file_path = params.get("file_path", "")
    target = _validate_path(file_path, sandbox_dir)
    if target is None:
        print(json.dumps({"status": "error", "error": "非法路径"}, ensure_ascii=False))
        return

    if not os.path.exists(target):
        print(json.dumps({"status": "error", "error": "文件不存在: %s" % file_path}, ensure_ascii=False))
        return

    if not target.lower().endswith(".docx"):
        print(json.dumps({"status": "error", "error": "只支持 .docx 文件"}, ensure_ascii=False))
        return

    mode = params.get("mode", MODE_PREVIEW)
    if mode not in _READ_MODES:
        print(json.dumps({"status": "error", "error": "未知读取模式: %s" % mode}, ensure_ascii=False))
        return

    handler = _READ_MODES[mode]
    result = handler(params, target)
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()


class DocReader:
    """Word 文档文本提取器（供 server.py 的 KB 文件解析和上传端点调用）"""

    def extract_text(self, file_path: str) -> str:
        """提取 docx 文件的全部纯文本（含表格）

        Args:
            file_path: .docx 文件的绝对路径

        Returns:
            拼接后的纯文本字符串，段落之间用换行分隔，
            表格每行用 " | " 分隔各单元格。
        """
        from docx import Document
        doc = Document(file_path)
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def count_images(self, file_path: str) -> int:
        """统计 docx 文件中的图片数量

        双重检测：
        1. doc.part.rels 中所有 image 类型的关联部件（常规嵌入图片）
        2. 文档正文中的 InlineShape 数量（粘贴/内嵌图片）

        Args:
            file_path: .docx 文件的绝对路径

        Returns:
            图片数量（0 表示无图片）
        """
        try:
            from docx import Document
            doc = Document(file_path)
            count = 0
            # 方法1：遍历关联部件
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    count += 1
            # 方法2：统计内嵌形状（InlineShape）— 微信粘贴图片等
            try:
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') or \
                           run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
                            count += 1
            except Exception:
                pass
            return count
        except Exception:
            return 0
